import json
import pandas as pd
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# --- Configuration ---
# Use Absolute Path logic to fix "Missing PDF" on Mac
BASE_DIR = os.path.dirname(os.path.abspath(__file__)) 

# UPDATE THIS PATH to your exact parquet/jsonl location
INPUT_REL_PATH = "output/agri-advisory-job_1767082735/parsed_results.parquet"

INPUT_PATH = os.path.abspath(INPUT_REL_PATH)
OUTPUT_DOCX = os.path.abspath("final_report_styled.docx")
OUTPUT_PDF = os.path.abspath("final_report_styled.pdf")

FONT_ENGLISH = 'Calibri'
FONT_HINDI = 'Noto Sans Devanagari' 

def contains_devanagari(text):
    return any('\u0900' <= c <= '\u097F' for c in text)

def clean_text(text):
    """
    Fixes double-escaped newlines and removes code block markers.
    """
    if not text: return "N/A"
    text = str(text)
    # Fix literal "\n" strings often found in JSON dumps
    text = text.replace('\\n', '\n').replace('<br>', '\n')
    # Remove wrapper artifacts
    text = text.replace("```json", "").replace("```markdown", "").replace("```", "")
    return text.strip()

def iter_entries(input_path):
    ext = os.path.splitext(input_path)[1].lower()
    if ext == ".jsonl":
        with open(input_path, "r", encoding="utf-8") as f:
            for line in f:
                yield json.loads(line)
    else:
        df = pd.read_parquet(input_path)
        for record in df.to_dict(orient="records"):
            yield record

def format_run(run, text, bold=False, italic=False, size=11, color=None):
    """Applies font, size, language settings, and bold/italic."""
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color: 
        run.font.color.rgb = color
    
    # Handle Hindi/Regional Fonts
    if contains_devanagari(text):
        run.font.name = FONT_HINDI
        run._element.rPr.rFonts.set(qn('w:cs'), FONT_HINDI)
    else:
        run.font.name = FONT_ENGLISH

def process_markdown_content(doc, raw_text, base_indent=0):
    """
    Parses Markdown-like text and applies native Word formatting.
    Handles:
    - Bullet points (* or -)
    - Bold text (**text**)
    - Headers (# or ##)
    """
    lines = clean_text(raw_text).split('\n')
    
    for line in lines:
        line = line.strip()
        if not line: continue

        # 1. Determine Paragraph Style
        style = 'Normal'
        indent = base_indent
        
        # Heading detection
        if line.startswith('# '):
            style = 'Heading 3' # Map # to H3 for sizing reasons
            line = line[2:]
        elif line.startswith('## '):
            style = 'Heading 4'
            line = line[3:]
        # Bullet detection
        elif line.startswith('* ') or line.startswith('- '):
            style = 'List Bullet'
            line = line[2:]
        elif line.startswith('• '):
            style = 'List Bullet'
            line = line[2:]
        
        # 2. Create Paragraph
        p = doc.add_paragraph(style=style)
        if indent > 0:
            p.paragraph_format.left_indent = Inches(indent)
        
        # 3. Parse Bold Markup (**text**)
        # This splits "A **bold** word" into ['A ', '**bold**', ' word']
        parts = re.split(r'(\*\*.*?\*\*)', line)
        
        for part in parts:
            if not part: continue
            
            is_bold_chunk = False
            text_chunk = part
            
            if part.startswith('**') and part.endswith('**'):
                is_bold_chunk = True
                text_chunk = part[2:-2] # Remove **
            
            # Add run to paragraph
            run = p.add_run(text_chunk)
            format_run(run, text_chunk, bold=is_bold_chunk, size=11)

def add_kv_section(doc, raw_text):
    """Formatted User Input Context."""
    lines = clean_text(raw_text).split('\n')
    for line in lines:
        line = line.strip()
        if not line: continue
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2) 
        
        if ':' in line:
            parts = line.split(':', 1)
            key, val = parts[0] + ":", parts[1]
            
            # Key (Bold, Dark Gray)
            run_k = p.add_run(key)
            format_run(run_k, key, bold=True, size=10, color=RGBColor(60, 60, 60))
            
            # Value (Normal)
            run_v = p.add_run(val)
            format_run(run_v, val, bold=False, size=10)
        else:
            run = p.add_run(line)
            format_run(run, line, size=10)

def add_entry(doc, obj, idx):
    # --- VISUAL SEPARATOR ---
    # Create a solid line using a paragraph border logic is hard in python-docx, 
    # so we use a visual text separator or rely on page breaks.
    
    # --- HEADER: Scenario ID ---
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    # Add colored background logic if needed, but let's stick to text for stability
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"SCENARIO RECORD #{idx} | ID: {obj.get('custom_id', 'N/A')}")
    format_run(run, run.text, bold=True, size=14, color=RGBColor(0, 51, 102)) # Navy Blue
    doc.add_paragraph() # Spacer


    # --- 1. SYSTEM INSTRUCTIONS (Footer/Meta) ---
    doc.add_paragraph("_" * 90).alignment = WD_ALIGN_PARAGRAPH.CENTER # Thin separator
    p_sys_head = doc.add_paragraph()
    format_run(p_sys_head.add_run("System Instructions Used:"), "Sys", bold=True, size=10, color=RGBColor(100, 100, 100))
    p_sys = doc.add_paragraph(clean_text(obj.get('system_instruction', '')))
    for run in p_sys.runs:
        format_run(run, run.text, size=10, italic=True, color=RGBColor(105,105,105))

    # --- 2. USER CONTEXT (Structured) ---
    h1 = doc.add_paragraph()
    format_run(h1.add_run("USER INPUT CONTEXT"), "USER INPUT", bold=True, size=10, color=RGBColor(100, 100, 100))
    # Apply shading or box logic here effectively creates a "card"
    add_kv_section(doc, obj.get('prompt', ''))
    doc.add_paragraph() 


    # --- 3. THOUGHTS (Internal Logic - Indented) ---
    thoughts = obj.get('thoughts', '')
    if thoughts:
        h3 = doc.add_paragraph()
        format_run(h3.add_run("INTERNAL REASONING (Chain of Thought)"), "CoT", bold=True, size=10, color=RGBColor(200, 100, 0)) # Orange
        
        # Pass indent=0.25 to indentation logic
        process_markdown_content(doc, thoughts, base_indent=0.25)
        

    # --- 4. ADVISORY (The Core Result) ---
    h4 = doc.add_paragraph()
    # Green header for the solution
    format_run(h4.add_run("GENERATED ADVISORY"), "ADVISORY", bold=True, size=10, color=RGBColor(0, 128, 0)) 
    
    # Use the intelligent Markdown parser here
    process_markdown_content(doc, obj.get('advisory', ''))
    doc.add_paragraph()


    # Start new page for next entry
    doc.add_page_break()

def create_docx(input_p, output_p):
    doc = Document()
    
    # Set Narrow Margins for better fit (0.5 inches)
    section = doc.sections[0]
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)

    print(f"Reading from: {input_p}")
    count = 0
    for idx, obj in enumerate(iter_entries(input_p), 1): 
        add_entry(doc, obj, idx)
        count = idx
        if idx % 100 == 0: print(f"Processed {idx} entries...")
            
    doc.save(output_p)
    print(f"DOCX Saved: {output_p} with {count} entries.")

def convert_to_pdf(docx_p, pdf_p):
    try:
        from docx2pdf import convert
        print(f"Converting to PDF... (This uses MS Word, please wait)")
        convert(docx_p, pdf_p)
        print(f"PDF Saved: {pdf_p}")
    except Exception as e:
        print(f"Conversion Failed: {e}")

if __name__ == "__main__":
    if not os.path.exists(INPUT_PATH):
        print(f"Error: Input file not found at {INPUT_PATH}")
    else:
        create_docx(INPUT_PATH, OUTPUT_DOCX)
        convert_to_pdf(OUTPUT_DOCX, OUTPUT_PDF)